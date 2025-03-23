"""
Document Parser Service

This module provides functionality for parsing different document formats,
extracting text content from PDF, DOCX, and TXT files.
"""

import io
from pathlib import Path
from typing import BinaryIO

import pypdf

class DocumentParser:
    """
    Service for parsing different document formats (PDF, DOCX, TXT) to extract text.
    
    Primarily focused on CV/resume parsing but can be used for other documents as well.
    """
    
    @staticmethod
    def parse_document(file_obj: BinaryIO, filename: str) -> str:
        """
        Parse a document file and extract its text content.
        
        Args:
            file_obj: File-like object containing the document
            filename: Original filename with extension
            
        Returns:
            Extracted text content as a string
            
        Raises:
            ValueError: If the file format is not supported or parsing fails
        """
        file_extension = Path(filename).suffix.lower()
        
        if file_extension == ".pdf":
            return DocumentParser._parse_pdf(file_obj)
        elif file_extension == ".docx":
            return DocumentParser._parse_docx(file_obj)
        elif file_extension == ".txt":
            return DocumentParser._parse_txt(file_obj)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats are PDF, DOCX, and TXT.")
    
    @staticmethod
    def _parse_pdf(file_obj: BinaryIO) -> str:
        """Parse a PDF file and extract text content."""
        try:
            # Create a PDF reader object
            pdf_reader = pypdf.PdfReader(file_obj)
            
            # Extract text from each page
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def _parse_docx(file_obj: BinaryIO) -> str:
        """Parse a DOCX file and extract text content."""
        try:
            # Check if python-docx is installed
            try:
                import docx
            except ImportError:
                raise ValueError("python-docx is not installed. Install it with 'pip install python-docx'")
            
            doc = docx.Document(file_obj)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return "\n".join(full_text).strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def _parse_txt(file_obj: BinaryIO) -> str:
        """Parse a TXT file and extract text content."""
        try:
            # Read the file content as text
            content = file_obj.read()
            
            # Decode if it's bytes
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='replace')
            
            return content.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse TXT: {str(e)}")
